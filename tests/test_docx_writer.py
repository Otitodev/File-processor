"""
Tests for src/docx_writer.py

Generates actual Word documents and inspects their content using python-docx.
No mocking required.
"""

import io
import pytest
from docx import Document

from src.docx_writer import generate_word_document
from src.report_analyzer import ReportSummary


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def make_summaries(n: int = 2) -> list:
    return [
        ReportSummary(
            report_index=i + 1,
            title=f"IME Report {i + 1}",
            start_page=i * 5 + 1,
            end_page=(i + 1) * 5,
            summary=f"First paragraph of report {i + 1}.\n\nSecond paragraph of report {i + 1}.",
        )
        for i in range(n)
    ]


def doc_text(result: bytes) -> str:
    """Extract all paragraph text from docx bytes."""
    doc = Document(io.BytesIO(result))
    return " ".join(p.text for p in doc.paragraphs)


# ---------------------------------------------------------------------------
# generate_word_document
# ---------------------------------------------------------------------------

class TestGenerateWordDocument:
    def test_returns_bytes(self):
        result = generate_word_document("Jane Smith", make_summaries())
        assert isinstance(result, bytes)
        assert len(result) > 0

    def test_output_is_valid_docx(self):
        result = generate_word_document("Jane Smith", make_summaries())
        doc = Document(io.BytesIO(result))  # must not raise
        assert doc is not None

    def test_contains_claimant_name(self):
        result = generate_word_document("Jane Smith", make_summaries())
        assert "Jane Smith" in doc_text(result)

    def test_contains_each_report_title(self):
        result = generate_word_document("Jane Smith", make_summaries(2))
        text = doc_text(result)
        assert "IME Report 1" in text
        assert "IME Report 2" in text

    def test_contains_page_range(self):
        summaries = [ReportSummary(1, "Report", 3, 7, "Summary.")]
        result = generate_word_document("Jane Smith", summaries)
        text = doc_text(result)
        assert "3" in text
        assert "7" in text

    def test_summary_split_on_double_newline(self):
        """Each double-newline-separated block should be its own paragraph."""
        summaries = [ReportSummary(1, "Report", 1, 5, "First para.\n\nSecond para.")]
        result = generate_word_document("Jane Smith", summaries)
        doc = Document(io.BytesIO(result))
        texts = [p.text for p in doc.paragraphs if p.text.strip()]
        assert any("First para." in t for t in texts)
        assert any("Second para." in t for t in texts)

    def test_empty_summaries_produces_valid_doc(self):
        result = generate_word_document("Jane Smith", [])
        doc = Document(io.BytesIO(result))
        assert "Jane Smith" in doc_text(result)

    def test_multiple_summaries_all_present(self):
        summaries = make_summaries(5)
        result = generate_word_document("Jane Smith", summaries)
        text = doc_text(result)
        for i in range(1, 6):
            assert f"IME Report {i}" in text

    def test_generated_heading_contains_title(self):
        summaries = [ReportSummary(1, "Neuropsychological Assessment", 1, 10, "Summary.")]
        result = generate_word_document("Test Claimant", summaries)
        doc = Document(io.BytesIO(result))
        heading_texts = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
        assert any("Neuropsychological Assessment" in t for t in heading_texts)
