"""
Word document writer for medical-legal report summaries.

Generates a formatted .docx file from a list of ReportSummary objects.
"""

import io
from datetime import date
from typing import List

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .report_analyzer import ReportSummary


def generate_word_document(claimant_name: str, summaries: List[ReportSummary]) -> bytes:
    """
    Build a Word document containing all report summaries.

    Args:
        claimant_name: The claimant's full name (used in the title).
        summaries:     Ordered list of ReportSummary objects to include.

    Returns:
        The .docx file contents as raw bytes, ready to write to disk or serve
        as a download.
    """
    doc = Document()

    # -----------------------------------------------------------------------
    # Title block
    # -----------------------------------------------------------------------
    title_para = doc.add_heading("Medical-Legal Summary", level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    claimant_para = doc.add_heading(f"Claimant: {claimant_name}", level=1)
    claimant_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    date_para = doc.add_paragraph(f"Generated: {date.today().strftime('%B %d, %Y')}")
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.runs[0].font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    doc.add_paragraph()  # spacer

    # -----------------------------------------------------------------------
    # One section per report
    # -----------------------------------------------------------------------
    for summary in summaries:
        # Report heading
        doc.add_heading(summary.title, level=2)

        # Page range sub-line in a slightly muted style
        page_para = doc.add_paragraph(f"Pages {summary.start_page}–{summary.end_page}")
        page_para.runs[0].font.size = Pt(10)
        page_para.runs[0].font.color.rgb = RGBColor(0x60, 0x60, 0x60)

        # Summary body — split on blank lines so each prose paragraph becomes
        # its own Word paragraph (mirrors the two-paragraph IME format, etc.)
        for block in summary.summary.split("\n\n"):
            block = block.strip()
            if block:
                doc.add_paragraph(block)

        doc.add_paragraph()  # spacer between reports

    # -----------------------------------------------------------------------
    # Serialize to bytes
    # -----------------------------------------------------------------------
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
