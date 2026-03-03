"""
Word document (.docx) text extractor.

Extracts text from a .docx file and splits it into virtual pages so it can
be fed into the same boundary-detection pipeline as PDF-derived text.
"""

import io
from typing import List, Union


def docx_to_page_texts(
    source: Union[str, bytes],
    chars_per_page: int = 3000,
) -> List[str]:
    """
    Extract text from a .docx file and split into virtual pages.

    Uses python-docx (already in requirements). Returns a list of text
    strings, each representing one virtual page (~3000 chars by default).
    Page numbers are 1-based and match the returned list index + 1.

    Args:
        source:         Path to a .docx file, or raw .docx bytes.
        chars_per_page: Target character count per virtual page. Paragraph
                        boundaries are respected — a paragraph is never split
                        across pages.

    Returns:
        List of text strings (one per virtual page). Never empty; if the
        document has no text, returns a list with one empty string.
    """
    from docx import Document  # python-docx

    if isinstance(source, bytes):
        doc = Document(io.BytesIO(source))
    else:
        doc = Document(source)

    # Collect all paragraph texts from the body and from table cells
    paragraphs: List[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    paragraphs.append(text)

    if not paragraphs:
        return [""]

    # Bin paragraphs into virtual pages without splitting any single paragraph
    pages: List[str] = []
    current_parts: List[str] = []
    current_len = 0

    for para in paragraphs:
        if current_len > 0 and current_len + len(para) > chars_per_page:
            pages.append("\n\n".join(current_parts))
            current_parts = []
            current_len = 0
        current_parts.append(para)
        current_len += len(para)

    if current_parts:
        pages.append("\n\n".join(current_parts))

    return pages
